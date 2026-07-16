"""
基于 V3.0 PRD 更新为 V4.0
新增内容：文本节点多状态、节点封面文案、标题双击编辑、添加节点默认展开、
工具箱页面、画布资产导入弹窗、参考图上传三选项、下载功能等
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from copy import deepcopy
import os

# 读取 V3.0 作为基础模板
src = 'PRD_无限画布创作工具_V3.0.docx'
doc = Document(src)

# ---- 修改版本号 ----
for para in doc.paragraphs:
    if 'PRD V3.0' in para.text:
        for run in para.runs:
            run.text = run.text.replace('V3.0', 'V4.0')
        break

# ---- 在版本历史表后追加新内容 ----
# 找到文件末尾位置，追加新章节

# 辅助函数
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_normal(doc, text):
    p = doc.add_paragraph(text, style='Normal')
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    return table

# ---- 追加 V4.0 更新日志 ----
doc.add_page_break()
add_heading(doc, '7. V4.0 更新日志（2026-07-16）', level=1)

add_heading(doc, '7.1 工作台页面（index.html）', level=2)
add_table(doc,
    ['功能项', '说明', '状态'],
    [
        ['SPA架构', '工作台/项目详情/工具箱页面切换', '✅'],
        ['步骤导航', '①资产管理 → ②分镜剧本 → ③视频预览', '✅'],
        ['创建/进入无限画布', '项目详情顶部入口按钮（Beta），跳转canvas.html', '✅'],
        ['任务中心按钮', '项目详情顶部右侧', '✅'],
    ]
)

add_heading(doc, '7.2 工具箱页面', level=2)
add_table(doc,
    ['功能项', '说明', '状态'],
    [
        ['配音创作', '功能卡片入口', '✅'],
        ['图片创作', '功能卡片入口', '✅'],
        ['视频创作', '功能卡片入口', '✅'],
        ['创建/进入无限画布', '功能卡片入口，跳转canvas.html', '✅'],
    ]
)

add_heading(doc, '7.3 项目详情页', level=2)
add_table(doc,
    ['功能项', '说明', '状态'],
    [
        ['资产类型垂直导航', '角色/场景/道具/通用素材 左侧导航', '✅'],
        ['角色列表', '搜索框+角色卡片（头像+名称+状态数）', '✅'],
        ['主视图/三视图切换', 'Tab切换，主视图大图预览', '✅'],
        ['图片预览+工具栏', '编辑/下载/全屏/列数 工具按钮', '✅'],
        ['信息栏', '来源/模型/时间/设为参考图', '✅'],
        ['生成记录', '标题+批量生成+缩略图列表', '✅'],
        ['AI生图配置面板', '模型/分辨率/数量/比例/参考图/提示词/确认生图', '✅'],
        ['参考图上传', 'hover展开：本地上传/资产库导入/画布资产导入', '✅'],
    ]
)

add_heading(doc, '7.4 画布资产导入弹窗', level=2)
add_table(doc,
    ['功能项', '说明', '状态'],
    [
        ['暗色主题', '#1e1e2e 背景', '✅'],
        ['标题', '画布资产导入', '✅'],
        ['缩放控制', '− 100% +', '✅'],
        ['筛选标签', '全部/角色生成/场景生成/道具生成/图片生成/分镜视频/视频合成', '✅'],
        ['卡片网格', '按日期分组，缩略图+复选框+名称+类型标签', '✅'],
        ['批量操作栏', '已选N项/下载/删除/使用/取消选择', '✅'],
        ['Tab筛选', '点击标签过滤对应类型卡片', '✅'],
    ]
)

add_heading(doc, '7.5 文本节点多状态重构', level=2)
add_normal(doc, '文本节点（TextNodePanel）重构为多状态交互组件，支持以下状态流转：')
add_table(doc,
    ['状态', '触发条件', '界面内容', '状态'],
    [
        ['编辑状态', '默认/添加节点后', '@提示+文本输入框+模型+⚡1生成', '✅'],
        ['生成中', '点击生成按钮', '加载动画+"AI 正在生成中..."', '✅'],
        ['生成结果', '3秒后自动完成', '编辑面板保持显示，结果通过setDesc显示在卡片内容区', '✅'],
    ]
)
add_normal(doc, '关键交互：')
add_bullet(doc, '添加节点后编辑面板默认展开（expanded=true）')
add_bullet(doc, '卡片内容区域始终显示：有内容显示文本，无内容显示空状态提示')
add_bullet(doc, '生成结果自动写入卡片内容区（setDesc）')
add_bullet(doc, '卡片顶部标题栏显示⬇下载按钮（仅有内容时），点击导出txt格式')
add_bullet(doc, '生成结果后编辑面板保持显示，可继续输入新内容再次生成')
add_bullet(doc, 'React.useEffect置于组件顶层，避免Hooks规则违反')

add_heading(doc, '7.6 节点封面文案统一', level=2)
add_table(doc,
    ['节点类型', '空状态封面文案', '图标', '状态'],
    [
        ['文本节点', '双击直接编写，或发起生成', '📄', '✅'],
        ['角色生成', '描述内容，开启角色创作', '👤', '✅'],
        ['场景生成', '描述内容，开启场景创作', '🏞️', '✅'],
        ['道具生成', '描述内容，开启道具创作', '💎', '✅'],
        ['图片生成', '描述内容，开启图片创作', '🖼️', '✅'],
    ]
)

add_heading(doc, '7.7 节点标题双击编辑', level=2)
add_bullet(doc, 'CustomNode（文本/分镜/合成）：标题从span改为input，支持双击编辑')
add_bullet(doc, 'CharacterNode（角色）：标题为input，支持编辑角色名称')
add_bullet(doc, 'VisualNode（场景/道具）：标题从span改为input，nodeLabel改为useState管理')
add_bullet(doc, 'ImageNode（图片）：标题支持编辑')

add_heading(doc, '7.8 节点底部信息栏调整', level=2)
add_normal(doc, '以下节点的底部信息栏已去掉：')
add_bullet(doc, '角色生成：去掉"基础形象/出现集数/语音/更多"信息栏')
add_bullet(doc, '图片生成：去掉"图片生成/模型/语音/更多"信息栏')
add_bullet(doc, '场景/道具生成：去掉"场景生成/待生成/更多"信息栏')
add_bullet(doc, '角色生成：图片编辑工具栏（选择框/圆形/画笔/三角形/撤销/重做/刷新/取消/擦除）已去掉')

# ---- 新增验收标准 ----
doc.add_page_break()
add_heading(doc, '8. V4.0 新增验收标准', level=1)

add_heading(doc, '8.1 工作台与工具箱', level=2)
add_table(doc,
    ['编号', '验收条件', 'EARS类型'],
    [
        ['AC-V4-01', 'When user clicks "工具箱" in nav, the system shall display toolbox page with 4 function cards', 'Event-driven'],
        ['AC-V4-02', 'When user clicks "创建/进入无限画布" card in toolbox, the system shall navigate to canvas.html', 'Event-driven'],
        ['AC-V4-03', 'The system shall display step navigation: 资产管理→分镜剧本→视频预览', 'Ubiquitous'],
        ['AC-V4-04', 'When user clicks "创建/进入无限画布" in project detail, the system shall navigate to canvas.html with project ID', 'Event-driven'],
    ]
)

add_heading(doc, '8.2 项目详情页', level=2)
add_table(doc,
    ['编号', '验收条件', 'EARS类型'],
    [
        ['AC-V4-05', 'The system shall display asset type navigation: 角色/场景/道具/通用素材', 'Ubiquitous'],
        ['AC-V4-06', 'The system shall display AI config panel with model/resolution/ratio/reference/prompt', 'Ubiquitous'],
        ['AC-V4-07', 'When user hovers over reference upload area, the system shall show 3 options: 本地上传/资产库导入/画布资产导入', 'Event-driven'],
        ['AC-V4-08', 'When user clicks "画布资产导入", the system shall show dark theme modal with filter tabs and card grid', 'Event-driven'],
    ]
)

add_heading(doc, '8.3 画布资产弹窗', level=2)
add_table(doc,
    ['编号', '验收条件', 'EARS类型'],
    [
        ['AC-V4-09', 'The system shall display asset cards in dark theme (#1e1e2e) grouped by date', 'Ubiquitous'],
        ['AC-V4-10', 'The system shall display filter tabs: 全部/角色生成/场景生成/道具生成/图片生成/分镜视频/视频合成', 'Ubiquitous'],
        ['AC-V4-11', 'When user clicks a filter tab, the system shall filter cards by type', 'Event-driven'],
        ['AC-V4-12', 'When user selects cards, the system shall show batch action bar with count and operations', 'Event-driven'],
        ['AC-V4-13', 'When user clicks "使用", the system shall import selected assets and close modal', 'Event-driven'],
    ]
)

add_heading(doc, '8.4 节点系统', level=2)
add_table(doc,
    ['编号', '验收条件', 'EARS类型'],
    [
        ['AC-V4-14', 'When user adds a new node, the system shall expand the editing panel by default', 'Event-driven'],
        ['AC-V4-15', 'When user double-clicks a node title, the system shall enable inline editing', 'Event-driven'],
        ['AC-V4-16', 'The system shall display empty state prompts: text→双击直接编写, character→描述内容开启角色创作, scene→描述内容开启场景创作, prop→描述内容开启道具创作, image→描述内容开启图片创作', 'Ubiquitous'],
        ['AC-V4-17', 'When text node has content (desc), the system shall display content text on the card', 'State-driven'],
        ['AC-V4-18', 'When text node has content, the system shall display ⬇ download button in the card header', 'State-driven'],
        ['AC-V4-19', 'When user clicks download, the system shall export content as .txt file', 'Event-driven'],
        ['AC-V4-20', 'The system shall not display bottom info bar on character/scene/prop/image nodes', 'Ubiquitous'],
    ]
)

add_heading(doc, '8.5 文本节点多状态', level=2)
add_table(doc,
    ['编号', '验收条件', 'EARS类型'],
    [
        ['AC-V4-21', 'When user adds a text node, the system shall display edit panel with @hint + textarea + model + ⚡1 generate button', 'Event-driven'],
        ['AC-V4-22', 'When user clicks "⚡1 生成", the system shall display loading animation with "AI 正在生成中..."', 'Event-driven'],
        ['AC-V4-23', 'When generation completes (3s), the system shall write result to desc and display on card, keeping edit panel visible', 'Event-driven'],
        ['AC-V4-24', 'While text node is in result state, the system shall keep @hint + textarea + model + ⚡1 generate visible', 'State-driven'],
        ['AC-V4-25', 'The system shall place React.useEffect at component top level, not inside conditional branches', 'Ubiquitous'],
    ]
)

# ---- 更新文本节点章节（3.1）----
add_heading(doc, '9. V4.0 文本节点详细说明（替换3.1节）', level=1)

add_heading(doc, '9.1 节点卡片', level=2)
add_bullet(doc, '标题栏：📜图标 + 可编辑标题（input，双击编辑）+ 下载按钮⬇（有内容时显示）+ 展开/收起箭头▼')
add_bullet(doc, '内容区域：有内容(desc)时显示文本内容（12px，#c1c2c5，maxHeight:160px，overflow:auto）')
add_bullet(doc, '空状态：📄图标 + "双击直接编写，或发起生成"')
add_bullet(doc, '下载功能：点击⬇按钮，导出desc为txt文件，文件名为节点标题.txt')

add_heading(doc, '9.2 编辑面板（展开后）', level=2)
add_bullet(doc, '@提示栏："使用 @ 指定引用素材用途（非必须）" + ✕关闭按钮')
add_bullet(doc, '文本输入区：多行textarea（minHeight:100px），placeholder提示创作需求示例')
add_bullet(doc, '底部栏：模型标签（🤖 Lux Gem 3.0）+ ⚡1生成按钮（渐变紫色，有内容时可点击）')

add_heading(doc, '9.3 生成中状态', level=2)
add_bullet(doc, '触发：点击⚡1生成按钮')
add_bullet(doc, '显示：40px旋转动画 + "AI 正在生成中..." + "画布文本生成任务"')
add_bullet(doc, '时长：3秒模拟（实际对接后端API）')

add_heading(doc, '9.4 生成结果状态', level=2)
add_bullet(doc, '自动完成：3秒后setGenResult写入结果，setDesc同步到卡片')
add_bullet(doc, '编辑面板保持显示：@提示+文本输入框+模型+⚡1生成（与初始状态一致）')
add_bullet(doc, '可继续生成：在输入框输入新内容，再次点击⚡1生成')
add_bullet(doc, '结果文本显示在卡片内容区域（不在面板内重复显示）')

# ---- 保存 ----
output = 'PRD_无限画布创作工具_V4.0.docx'
doc.save(output)
print(f'✅ 已生成: {output}')
print(f'文件大小: {os.path.getsize(output)} bytes')
